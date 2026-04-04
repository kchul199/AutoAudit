"""
CP1 — 도메인 문서 로더
────────────────────────────────────────────────────
지원 형식: PDF, TXT, HTML, DOCX, URL(크롤링)
문서 타입 자동 감지: FAQ | 매뉴얼 | 웹사이트

출력: List[RawDocument]
  RawDocument.doc_id      : 고유 문서 ID
  RawDocument.subscriber  : 가입자명
  RawDocument.doc_type    : "faq" | "manual" | "website"
  RawDocument.title       : 문서 제목
  RawDocument.content     : 전체 텍스트
  RawDocument.source      : 원본 파일 경로 또는 URL
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional
from urllib.parse import urljoin, urlparse

from app.utils.logger import get_logger

logger = get_logger(__name__)


# ── 데이터 클래스 ─────────────────────────────────────────────────

@dataclass
class RawDocument:
    doc_id: str
    subscriber: str
    doc_type: str        # "faq" | "manual" | "website"
    title: str
    content: str
    source: str
    metadata: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "doc_id":     self.doc_id,
            "subscriber": self.subscriber,
            "doc_type":   self.doc_type,
            "title":      self.title,
            "content":    self.content[:200] + "..." if len(self.content) > 200 else self.content,
            "source":     self.source,
            "char_count": len(self.content),
            "metadata":   self.metadata,
        }


# ── 문서 타입 자동 감지 ──────────────────────────────────────────

_FAQ_KEYWORDS    = ["q:", "q.", "질문", "답변", "faq", "자주 묻는", "a:", "a."]
_MANUAL_KEYWORDS = ["목차", "chapter", "제", "항", "1.", "1.1", "절차", "주의", "단계",
                    "매뉴얼", "가이드", "설명서", "안내서"]
_WEB_KEYWORDS    = ["http", "www", "href", "html", "<div", "<p>", "홈페이지", "웹사이트"]


def detect_doc_type(text: str, filename: str = "") -> str:
    """텍스트·파일명을 분석하여 문서 타입 추론"""
    text_lower = text[:3000].lower()
    fname_lower = filename.lower()

    # 파일명 힌트 우선
    if any(k in fname_lower for k in ["faq", "자주묻는", "q&a", "qna"]):
        return "faq"
    if any(k in fname_lower for k in ["manual", "매뉴얼", "guide", "가이드", "안내서"]):
        return "manual"
    if any(k in fname_lower for k in ["web", "site", "홈페이지", "html"]):
        return "website"

    # 텍스트 패턴 분석
    faq_score    = sum(1 for k in _FAQ_KEYWORDS    if k in text_lower)
    manual_score = sum(1 for k in _MANUAL_KEYWORDS if k in text_lower)
    web_score    = sum(1 for k in _WEB_KEYWORDS    if k in text_lower)

    scores = {"faq": faq_score, "manual": manual_score, "website": web_score}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "manual"


# ── 문서 로더 클래스 ─────────────────────────────────────────────

class DocLoader:
    """
    가입자 도메인 문서를 로드하여 RawDocument 목록으로 변환.

    사용 예:
        loader = DocLoader(subscriber="한국통신")
        docs = loader.load_directory("data/docs/한국통신/")
        docs += loader.load_url("https://kt.com/support/faq")
    """

    def __init__(self, subscriber: str, doc_type_override: str = None):
        self.subscriber = subscriber
        self.doc_type_override = doc_type_override  # None이면 자동 감지

    # ── Public API ───────────────────────────────────────────────

    def load_file(self, file_path: str) -> RawDocument:
        """단일 파일 로드"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"파일 없음: {file_path}")

        ext = path.suffix.lower()
        logger.info(f"[CP1] 문서 로드: {path.name} ({ext})")

        if ext == ".pdf":
            content = self._load_pdf(path)
        elif ext in (".docx", ".doc"):
            content = self._load_docx(path)
        elif ext in (".html", ".htm"):
            content = self._load_html(path.read_text(encoding="utf-8", errors="replace"))
        else:
            content = path.read_text(encoding="utf-8", errors="replace")

        title = path.stem
        doc_type = self.doc_type_override or detect_doc_type(content, path.name)

        return RawDocument(
            doc_id=self._gen_id(file_path),
            subscriber=self.subscriber,
            doc_type=doc_type,
            title=title,
            content=self._clean_text(content),
            source=str(path.resolve()),
            metadata={"filename": path.name, "ext": ext},
        )

    def load_directory(self, dir_path: str) -> List[RawDocument]:
        """디렉토리 내 지원 파일 전체 로드"""
        exts = {".pdf", ".txt", ".html", ".htm", ".docx", ".doc", ".md"}
        files = [f for f in Path(dir_path).rglob("*") if f.suffix.lower() in exts]
        logger.info(f"[CP1] 디렉토리 문서 로드: {dir_path} ({len(files)}개 파일)")

        docs: List[RawDocument] = []
        for f in files:
            try:
                docs.append(self.load_file(str(f)))
            except Exception as e:
                logger.error(f"[CP1] 문서 로드 오류 {f.name}: {e}")
        return docs

    def load_url(self, url: str, max_pages: int = 5, depth: int = 1) -> List[RawDocument]:
        """
        URL 크롤링하여 웹 페이지 로드.
        depth: 링크 재귀 탐색 깊이 (1 = 시작 페이지만)
        """
        try:
            import requests
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("[CP1] requests/beautifulsoup4 미설치: pip install requests beautifulsoup4")
            return []

        visited: set = set()
        docs: List[RawDocument] = []
        queue = [(url, 0)]

        while queue and len(docs) < max_pages:
            current_url, current_depth = queue.pop(0)
            if current_url in visited:
                continue
            visited.add(current_url)

            try:
                resp = requests.get(current_url, timeout=15,
                                    headers={"User-Agent": "AutoAudit/1.0"})
                resp.raise_for_status()
                soup = BeautifulSoup(resp.text, "html.parser")
                content = self._load_html(resp.text)
                title = soup.title.string.strip() if soup.title else current_url

                doc_type = self.doc_type_override or detect_doc_type(content, current_url)
                docs.append(RawDocument(
                    doc_id=self._gen_id(current_url),
                    subscriber=self.subscriber,
                    doc_type=doc_type,
                    title=title,
                    content=self._clean_text(content),
                    source=current_url,
                    metadata={"url": current_url, "depth": current_depth},
                ))
                logger.info(f"[CP1] 크롤링: {current_url} → {doc_type}")

                # 내부 링크 수집 (depth 제한 내)
                if current_depth < depth:
                    base = f"{urlparse(current_url).scheme}://{urlparse(current_url).netloc}"
                    for a in soup.find_all("a", href=True)[:20]:
                        href = urljoin(base, a["href"])
                        if href.startswith(base) and href not in visited:
                            queue.append((href, current_depth + 1))

            except Exception as e:
                logger.warning(f"[CP1] 크롤링 실패 {current_url}: {e}")

        return docs

    def load_text(self, text: str, title: str = "inline", doc_type: str = None) -> RawDocument:
        """텍스트 직접 로드 (테스트·UI용)"""
        detected = doc_type or self.doc_type_override or detect_doc_type(text, title)
        return RawDocument(
            doc_id=self._gen_id(text[:200]),
            subscriber=self.subscriber,
            doc_type=detected,
            title=title,
            content=self._clean_text(text),
            source="inline",
        )

    # ── 내부 파서 ─────────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            logger.error("[CP1] pypdf 미설치: pip install pypdf")
            return ""
        except Exception as e:
            logger.error(f"[CP1] PDF 파싱 오류: {e}")
            return ""

    def _load_docx(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except ImportError:
            logger.error("[CP1] python-docx 미설치: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"[CP1] DOCX 파싱 오류: {e}")
            return ""

    @staticmethod
    def _load_html(html: str) -> str:
        """HTML에서 텍스트 추출 (BeautifulSoup 우선, 없으면 regex fallback)"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            # 스크립트/스타일 제거
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            return soup.get_text(separator="\n")
        except ImportError:
            # 간단 fallback
            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"&[a-z]+;", " ", text)
            return text

    @staticmethod
    def _clean_text(text: str) -> str:
        """불필요한 공백·특수문자 정규화"""
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r" {2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _gen_id(content: str) -> str:
        return hashlib.md5(content.encode("utf-8")).hexdigest()[:12]
